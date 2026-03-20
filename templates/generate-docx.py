#!/usr/bin/env python3
"""
企业实施方案 Docx 生成器
将 Markdown 格式的方案转换为符合规范的 Word 文档

用法:
    python generate-docx.py input.md output.docx --title "方案标题"
    python generate-docx.py input.md output.docx --title "标题" --number "编号" --author "张三"
"""

import sys
import re
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 常量 ──────────────────────────────────────────────
FONT_CN = '宋体'
FONT_EN = 'Times New Roman'
TITLE_SIZE = Pt(22)
H1_SIZE = Pt(12)
H2_SIZE = Pt(12)
H3_SIZE = Pt(12)
BODY_SIZE = Pt(10.5)
LINE_SPACING = Pt(22)
FIRST_INDENT = Cm(0.74)
TABLE_FONT_SIZE = Pt(10)
HEADER_SHADING = 'D9E2F3'

# 页面：A4，四边距 1.5cm
PAGE_W = Cm(21)
PAGE_H = Cm(29.7)
MARGIN = Cm(1.5)
# 内容区宽度 = 21 - 1.5*2 = 18cm
CONTENT_W = Cm(18)


def parse_args():
    p = argparse.ArgumentParser(description='Markdown 方案 → Docx')
    p.add_argument('input', help='Markdown 文件路径（- 为 stdin）')
    p.add_argument('output', help='输出 Docx 路径')
    p.add_argument('--title', default='实施方案', help='方案标题')
    p.add_argument('--number', default='', help='文档编号')
    p.add_argument('--author', default='', help='编制人')
    p.add_argument('--reviewer', default='', help='审核人')
    p.add_argument('--approver', default='', help='批准人')
    p.add_argument('--date', default='', help='生效日期')
    return p.parse_args()


# ── 字体工具 ──────────────────────────────────────────
def set_font(run, size, bold=False, color=None):
    run.font.size = size
    run.font.bold = bold
    run.font.name = FONT_EN
    if color:
        run.font.color.rgb = color
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_EN}" '
            f'w:hAnsi="{FONT_EN}" w:eastAsia="{FONT_CN}"/>'
        )
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), FONT_CN)
        rFonts.set(qn('w:ascii'), FONT_EN)
        rFonts.set(qn('w:hAnsi'), FONT_EN)


def add_run_with_highlight(paragraph, text, size, bold=False):
    parts = re.split(r'(【待确认[^】]*】)', text)
    for part in parts:
        if part.startswith('【待确认'):
            run = paragraph.add_run(part)
            set_font(run, size, bold=False)
            rPr = run._element.get_or_add_rPr()
            highlight = parse_xml(f'<w:highlight {nsdecls("w")} w:val="yellow"/>')
            rPr.append(highlight)
        else:
            bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
            for bp in bold_parts:
                if bp.startswith('**') and bp.endswith('**'):
                    run = paragraph.add_run(bp[2:-2])
                    set_font(run, size, bold=True)
                elif bp:
                    run = paragraph.add_run(bp)
                    set_font(run, size, bold=bold)


def _remove_cell_borders(cell):
    """移除单元格所有边框"""
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def _add_page_field(paragraph, field_code=' PAGE '):
    """在段落中插入域代码（如页码）"""
    run = paragraph.add_run()
    run._element.append(parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'
    ))
    run2 = paragraph.add_run()
    run2._element.append(parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">{field_code}</w:instrText>'
    ))
    run3 = paragraph.add_run()
    run3._element.append(parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'
    ))


# ── 页面设置 ──────────────────────────────────────────
def setup_page(doc):
    """A4 纸张，四边距 1.5cm"""
    section = doc.sections[0]
    section.page_width = PAGE_W
    section.page_height = PAGE_H
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)
    return section


def add_page_borders(section, sz=4, space=24, color='auto'):
    """页面边框，offsetFrom=page，space=距页边缘pt"""
    sectPr = section._sectPr
    pg_borders = parse_xml(
        f'<w:pgBorders {nsdecls("w")} w:offsetFrom="page">'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="{space}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="{space}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="{space}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="{space}" w:color="{color}"/>'
        f'</w:pgBorders>'
    )
    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is not None:
        pgSz.addnext(pg_borders)
    else:
        sectPr.append(pg_borders)


# ── 页眉（结构化：Logo | 标题 | 编号+页码）──────────────
def add_header_footer(doc, title, number):
    section = doc.sections[0]

    header = section.header
    header.is_linked_to_previous = False
    # 清除默认段落
    for p in header.paragraphs:
        p.clear()

    # 创建 3 列表格作为页眉栏
    h_table = header.add_table(rows=1, cols=3, width=CONTENT_W)
    h_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    h_table.autofit = False

    # 设置列宽：Logo 4cm | 标题 10cm | 编号+页码 4cm
    for row in h_table.rows:
        row.height = Cm(1.2)
        cells = row.cells
        cells[0].width = Cm(4)
        cells[1].width = Cm(10)
        cells[2].width = Cm(4)

    # Col 0: Logo 占位
    c0 = h_table.cell(0, 0)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p0.add_run('[Logo]')
    set_font(run, Pt(9), color=RGBColor(0x99, 0x99, 0x99))

    # Col 1: 标题
    c1 = h_table.cell(0, 1)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p1.add_run(title)
    set_font(run, Pt(9), bold=True)

    # Col 2: 编号 + 页码
    c2 = h_table.cell(0, 2)
    c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if number:
        run = p2.add_run(f'{number}  ')
        set_font(run, Pt(8))
    run = p2.add_run('第')
    set_font(run, Pt(8))
    _add_page_field(p2)
    run = p2.add_run('页')
    set_font(run, Pt(8))

    # 移除表格内部边框，只保留底部边框线
    for cell in h_table.row_cells(0):
        tcPr = cell._element.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)

    # 页脚 — 居中页码
    footer = section.footer
    footer.is_linked_to_previous = False
    pf = footer.paragraphs[0]
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_field(pf)


# ── 封面（封面 + 审批同页）────────────────────────────
def add_cover_page(doc, title, number, author='', reviewer='',
                   approver='', date=''):
    # 留白
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # 大图占位区
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run('[公司形象图 — 请替换]')
    set_font(run, Pt(16), color=RGBColor(0x99, 0x99, 0x99))

    # 方案标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(48)
    run = p.add_run(title)
    set_font(run, TITLE_SIZE, bold=True)

    # 审批信息表（无边框 2列4行）
    approval_table = doc.add_table(rows=4, cols=2)
    approval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    approval_table.autofit = False
    for row in approval_table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(6)

    fields = [
        ('编  制：', author or '________________'),
        ('审  核：', reviewer or '________________'),
        ('批  准：', approver or '________________'),
        ('生效日期：', date or '______年____月____日'),
    ]
    for i, (label, value) in enumerate(fields):
        for j in range(2):
            _remove_cell_borders(approval_table.cell(i, j))
        # label
        c_label = approval_table.cell(i, 0)
        p = c_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(label)
        set_font(run, Pt(12))
        # value
        c_val = approval_table.cell(i, 1)
        p = c_val.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(value)
        set_font(run, Pt(12))

    doc.add_page_break()


# ── Markdown 解析 ─────────────────────────────────────
def parse_markdown(text):
    blocks = []
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if re.match(r'^-{3,}$', line.strip()):
            blocks.append({'type': 'separator'})
            i += 1
            continue
        m = re.match(r'^(#{1,4})\s+(.+)$', line)
        if m:
            blocks.append({
                'type': 'heading', 'level': len(m.group(1)),
                'text': m.group(2).strip(),
            })
            i += 1
            continue
        if '|' in line and i + 1 < len(lines) and re.search(r'-{2,}', lines[i + 1]):
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            rows = []
            for tl in table_lines:
                if re.match(r'^[\s|:\-]+$', tl):
                    continue
                cells = [c.strip() for c in tl.strip().strip('|').split('|')]
                if cells:
                    rows.append(cells)
            if rows:
                blocks.append({'type': 'table', 'rows': rows})
            continue
        lm = re.match(
            r'^(\s*)([-*]|\d+[.、]|[①②③④⑤⑥⑦⑧⑨⑩])\s*(.+)$', line
        )
        if lm:
            blocks.append({
                'type': 'list_item', 'indent': len(lm.group(1)),
                'marker': lm.group(2), 'text': lm.group(3),
            })
            i += 1
            continue
        para_lines = [line.strip()]
        i += 1
        while (i < len(lines) and lines[i].strip()
               and not re.match(r'^#{1,4}\s', lines[i])
               and '|' not in lines[i]
               and not re.match(r'^\s*([-*]|\d+[.、]|[①-⑩])\s', lines[i])
               and not re.match(r'^-{3,}$', lines[i].strip())):
            para_lines.append(lines[i].strip())
            i += 1
        blocks.append({'type': 'paragraph', 'text': ''.join(para_lines)})
    return blocks


# ── 内容写入 ──────────────────────────────────────────
def add_content(doc, blocks):
    for block in blocks:
        btype = block['type']
        if btype == 'separator':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            continue
        if btype == 'heading':
            level = block['level']
            p = doc.add_paragraph()
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(block['text'])
                set_font(run, TITLE_SIZE, bold=True)
            elif level == 2:
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                add_run_with_highlight(p, block['text'], H1_SIZE, bold=True)
            elif level == 3:
                p.paragraph_format.left_indent = Cm(0.74)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
                add_run_with_highlight(p, block['text'], H2_SIZE, bold=True)
            elif level == 4:
                p.paragraph_format.left_indent = Cm(1.48)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(3)
                add_run_with_highlight(p, block['text'], H3_SIZE, bold=True)
            continue
        if btype == 'paragraph':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = FIRST_INDENT
            p.paragraph_format.line_spacing = LINE_SPACING
            add_run_with_highlight(p, block['text'], BODY_SIZE)
            continue
        if btype == 'list_item':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            indent_level = block['indent'] // 2
            p.paragraph_format.left_indent = Cm(0.74 * (indent_level + 1))
            p.paragraph_format.line_spacing = LINE_SPACING
            text = f"{block['marker']} {block['text']}"
            add_run_with_highlight(p, text, BODY_SIZE)
            continue
        if btype == 'table':
            add_table_block(doc, block['rows'])


def add_table_block(doc, rows):
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = 'Table Grid'

    # 表格撑满内容区宽度
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    existing = tblPr.find(qn('w:tblW'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblW)

    for i, row_data in enumerate(rows):
        for j in range(num_cols):
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_text = row_data[j] if j < len(row_data) else ''
            if i == 0:
                run = p.add_run(cell_text)
                set_font(run, TABLE_FONT_SIZE, bold=True)
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{HEADER_SHADING}" w:val="clear"/>'
                )
                cell._element.get_or_add_tcPr().append(shading)
            else:
                add_run_with_highlight(p, cell_text, TABLE_FONT_SIZE)

    doc.add_paragraph()


# ── 主函数 ────────────────────────────────────────────
def main():
    args = parse_args()

    if args.input == '-':
        md_text = sys.stdin.read()
    else:
        md_text = Path(args.input).read_text(encoding='utf-8')

    doc = Document()

    # 页面设置：A4，1.5cm 四边距，页面边框
    section = setup_page(doc)
    add_page_borders(section)

    # 页眉（Logo | 标题 | 编号+页码）+ 页脚（页码）
    add_header_footer(doc, args.title, args.number)

    # 封面 + 审批（同一页）
    add_cover_page(doc, args.title, args.number,
                   args.author, args.reviewer, args.approver, args.date)

    # 正文
    blocks = parse_markdown(md_text)
    add_content(doc, blocks)

    out = Path(args.output)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f'已生成: {out}')


if __name__ == '__main__':
    main()
